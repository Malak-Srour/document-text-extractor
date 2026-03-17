"""
Document Text Extractor
A Gradio-based application that extracts text from PDF and Word documents.
"""

import gradio as gr
import pdfplumber
from docx import Document
import os
from typing import Dict, List, Tuple


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
                else:
                    text_content.append(f"--- Page {page_num} ---\n[No text content found on this page]")
        return "\n\n".join(text_content) if text_content else "No text could be extracted from this PDF."
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        doc = Document(file_path)
        text_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content) if text_content else "No text could be extracted from this Word document."
    except Exception as e:
        return f"Error extracting text from Word document: {str(e)}"


def extract_text(file_path: str) -> str:
    """Extract text from a file based on its extension."""
    if not file_path:
        return "No file provided."
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension in [".docx", ".doc"]:
        if file_extension == ".doc":
            return "Note: .doc files (old Word format) are not fully supported. Please convert to .docx format.\n\nAttempting extraction anyway...\n" + extract_text_from_docx(file_path)
        return extract_text_from_docx(file_path)
    else:
        return f"Unsupported file format: {file_extension}. Please upload PDF or Word (.docx) files."


def process_files(files: List[str]) -> Tuple[gr.Dropdown, str, Dict[str, str]]:
    """Process uploaded files and return extracted text."""
    if not files:
        return (
            gr.Dropdown(choices=[], value=None, visible=False),
            "Please upload one or more PDF or Word documents.",
            {}
        )
    
    results = {}
    file_names = []
    
    for file_path in files:
        file_name = os.path.basename(file_path)
        file_names.append(file_name)
        extracted_text = extract_text(file_path)
        results[file_name] = extracted_text
    
    if len(files) == 1:
        # Single file - show directly without dropdown
        return (
            gr.Dropdown(choices=[], value=None, visible=False),
            f"📄 **{file_names[0]}**\n\n{results[file_names[0]]}",
            results
        )
    else:
        # Multiple files - show dropdown
        first_file = file_names[0]
        return (
            gr.Dropdown(
                choices=file_names,
                value=first_file,
                visible=True,
                label="Select a file to view"
            ),
            f"📄 **{first_file}**\n\n{results[first_file]}",
            results
        )


def update_display(selected_file: str, results: Dict[str, str]) -> str:
    """Update the text display when a file is selected from dropdown."""
    if not selected_file or not results:
        return "No file selected."
    return f"📄 **{selected_file}**\n\n{results.get(selected_file, 'No content available.')}"


# Create the Gradio interface
with gr.Blocks(
    title="Document Text Extractor",
    theme=gr.themes.Soft(
        primary_hue="blue",
        secondary_hue="gray",
    )
) as app:
    # Store results in state
    results_state = gr.State({})
    
    gr.Markdown(
        """
        # 📑 Document Text Extractor
        
        Upload PDF or Word documents to extract their text content.
        
        **Supported formats:** PDF (.pdf), Word (.docx)
        """
    )
    
    with gr.Row():
        with gr.Column(scale=1):
            file_input = gr.File(
                label="Upload Documents",
                file_count="multiple",
                file_types=[".pdf", ".docx", ".doc"],
                type="filepath"
            )
            
            extract_btn = gr.Button(
                "🔍 Extract Text",
                variant="primary",
                size="lg"
            )
            
            gr.Markdown(
                """
                ### Instructions
                1. Click "Upload Documents" or drag and drop files
                2. Select one or more PDF/Word files
                3. Click "Extract Text" to process
                4. View extracted text below
                5. Use dropdown to switch between files (if multiple)
                """
            )
    
    with gr.Row():
        with gr.Column():
            file_dropdown = gr.Dropdown(
                choices=[],
                label="Select a file to view",
                visible=False,
                interactive=True
            )
    
    with gr.Row():
        with gr.Column():
            output_text = gr.Markdown(
                value="*Upload documents and click 'Extract Text' to see results here.*",
                label="Extracted Text"
            )
    
    # Event handlers
    extract_btn.click(
        fn=process_files,
        inputs=[file_input],
        outputs=[file_dropdown, output_text, results_state]
    )
    
    file_dropdown.change(
        fn=update_display,
        inputs=[file_dropdown, results_state],
        outputs=[output_text]
    )
    
    # Example section
    gr.Markdown(
        """
        ---
        ### About
        
        This tool uses:
        - **pdfplumber** for PDF text extraction
        - **python-docx** for Word document processing
        
        The extracted text preserves the document structure as much as possible,
        including page breaks for PDFs and paragraph formatting for Word documents.
        """
    )


if __name__ == "__main__":
    app.launch()
