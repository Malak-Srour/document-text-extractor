"""
Document Text Extractor
A Flask-based application that extracts text from PDF and Word documents.
"""

from flask import Flask, render_template_string, request, jsonify
import pdfplumber
from docx import Document
import os
import tempfile
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pdfplumber."""
    text_content = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_content.append(f"--- Page {page_num} ---\n{page_text}")
                else:
                    text_content.append(f"--- Page {page_num} ---\n[No text content found on this page]")
        return "\n\n".join(text_content) if text_content else "No text could be extracted from this PDF."
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a Word document using python-docx."""
    try:
        doc = Document(file_path)
        text_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content) if text_content else "No text could be extracted from this Word document."
    except Exception as e:
        return f"Error extracting text from Word document: {str(e)}"


def extract_text(file_path: str, filename: str) -> str:
    """Extract text from a file based on its extension."""
    file_extension = os.path.splitext(filename)[1].lower()
    
    if file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension in [".docx", ".doc"]:
        if file_extension == ".doc":
            return "Note: .doc files (old Word format) are not fully supported. Please convert to .docx format.\n\nAttempting extraction anyway...\n" + extract_text_from_docx(file_path)
        return extract_text_from_docx(file_path)
    else:
        return f"Unsupported file format: {file_extension}. Please upload PDF or Word (.docx) files."


HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Document Text Extractor</title>
    <style>
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }

        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif;
            background-color: #f8fafc;
            color: #1e293b;
            min-height: 100vh;
            line-height: 1.6;
        }

        .container {
            max-width: 900px;
            margin: 0 auto;
            padding: 2rem 1rem;
        }

        header {
            text-align: center;
            margin-bottom: 2rem;
        }

        h1 {
            font-size: 2rem;
            font-weight: 700;
            color: #0f172a;
            margin-bottom: 0.5rem;
        }

        .subtitle {
            color: #64748b;
            font-size: 1rem;
        }

        .upload-section {
            background: #ffffff;
            border-radius: 12px;
            padding: 2rem;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            margin-bottom: 1.5rem;
        }

        .drop-zone {
            border: 2px dashed #cbd5e1;
            border-radius: 8px;
            padding: 3rem 2rem;
            text-align: center;
            cursor: pointer;
            transition: all 0.2s ease;
            background: #f8fafc;
        }

        .drop-zone:hover,
        .drop-zone.drag-over {
            border-color: #3b82f6;
            background: #eff6ff;
        }

        .drop-zone-icon {
            font-size: 3rem;
            margin-bottom: 1rem;
            color: #94a3b8;
        }

        .drop-zone-text {
            color: #475569;
            margin-bottom: 0.5rem;
        }

        .drop-zone-hint {
            color: #94a3b8;
            font-size: 0.875rem;
        }

        .file-input {
            display: none;
        }

        .file-list {
            margin-top: 1.5rem;
        }

        .file-item {
            display: flex;
            align-items: center;
            justify-content: space-between;
            padding: 0.75rem 1rem;
            background: #f1f5f9;
            border-radius: 6px;
            margin-bottom: 0.5rem;
        }

        .file-item-name {
            font-size: 0.875rem;
            color: #334155;
            display: flex;
            align-items: center;
            gap: 0.5rem;
        }

        .file-item-icon {
            color: #3b82f6;
        }

        .file-item-remove {
            background: none;
            border: none;
            color: #94a3b8;
            cursor: pointer;
            padding: 0.25rem;
            font-size: 1.25rem;
            line-height: 1;
            transition: color 0.2s;
        }

        .file-item-remove:hover {
            color: #ef4444;
        }

        .btn {
            display: inline-flex;
            align-items: center;
            justify-content: center;
            gap: 0.5rem;
            padding: 0.75rem 1.5rem;
            font-size: 1rem;
            font-weight: 500;
            border: none;
            border-radius: 8px;
            cursor: pointer;
            transition: all 0.2s ease;
            width: 100%;
            margin-top: 1rem;
        }

        .btn-primary {
            background: #3b82f6;
            color: white;
        }

        .btn-primary:hover:not(:disabled) {
            background: #2563eb;
        }

        .btn-primary:disabled {
            background: #94a3b8;
            cursor: not-allowed;
        }

        .results-section {
            background: #ffffff;
            border-radius: 12px;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
            overflow: hidden;
        }

        .results-header {
            padding: 1rem 1.5rem;
            border-bottom: 1px solid #e2e8f0;
            display: flex;
            align-items: center;
            gap: 1rem;
        }

        .results-title {
            font-weight: 600;
            color: #0f172a;
        }

        .file-selector {
            flex: 1;
            max-width: 300px;
        }

        .file-selector select {
            width: 100%;
            padding: 0.5rem 0.75rem;
            font-size: 0.875rem;
            border: 1px solid #e2e8f0;
            border-radius: 6px;
            background: white;
            color: #334155;
            cursor: pointer;
        }

        .file-selector select:focus {
            outline: none;
            border-color: #3b82f6;
            box-shadow: 0 0 0 3px rgba(59, 130, 246, 0.1);
        }

        .results-content {
            padding: 1.5rem;
            max-height: 500px;
            overflow-y: auto;
        }

        .text-output {
            font-family: 'SF Mono', Monaco, 'Cascadia Code', monospace;
            font-size: 0.875rem;
            line-height: 1.7;
            color: #334155;
            white-space: pre-wrap;
            word-wrap: break-word;
        }

        .placeholder {
            text-align: center;
            padding: 3rem;
            color: #94a3b8;
        }

        .loading {
            display: flex;
            align-items: center;
            justify-content: center;
            gap: 0.75rem;
            padding: 2rem;
            color: #64748b;
        }

        .spinner {
            width: 24px;
            height: 24px;
            border: 3px solid #e2e8f0;
            border-top-color: #3b82f6;
            border-radius: 50%;
            animation: spin 0.8s linear infinite;
        }

        @keyframes spin {
            to { transform: rotate(360deg); }
        }

        .hidden {
            display: none;
        }

        .error {
            background: #fef2f2;
            color: #dc2626;
            padding: 1rem;
            border-radius: 6px;
            margin-bottom: 1rem;
        }

        footer {
            text-align: center;
            padding: 2rem;
            color: #94a3b8;
            font-size: 0.875rem;
        }

        @media (max-width: 640px) {
            .container {
                padding: 1rem;
            }

            h1 {
                font-size: 1.5rem;
            }

            .upload-section {
                padding: 1.5rem;
            }

            .drop-zone {
                padding: 2rem 1rem;
            }

            .results-header {
                flex-direction: column;
                align-items: flex-start;
            }

            .file-selector {
                width: 100%;
                max-width: none;
            }
        }
    </style>
</head>
<body>
    <div class="container">
        <header>
            <h1>Document Text Extractor</h1>
            <p class="subtitle">Upload PDF or Word documents to extract their text content</p>
        </header>

        <section class="upload-section">
            <div class="drop-zone" id="dropZone">
                <div class="drop-zone-icon">
                    <svg width="48" height="48" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="1.5" d="M7 16a4 4 0 01-.88-7.903A5 5 0 1115.9 6L16 6a5 5 0 011 9.9M15 13l-3-3m0 0l-3 3m3-3v12"/>
                    </svg>
                </div>
                <p class="drop-zone-text">Drag and drop files here, or click to browse</p>
                <p class="drop-zone-hint">Supports PDF (.pdf) and Word (.docx) files</p>
            </div>
            <input type="file" id="fileInput" class="file-input" multiple accept=".pdf,.docx,.doc">
            
            <div class="file-list" id="fileList"></div>
            
            <button class="btn btn-primary" id="extractBtn" disabled>
                <svg width="20" height="20" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12h6m-6 4h6m2 5H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z"/>
                </svg>
                Extract Text
            </button>
        </section>

        <section class="results-section" id="resultsSection">
            <div class="results-header">
                <span class="results-title">Extracted Text</span>
                <div class="file-selector hidden" id="fileSelector">
                    <select id="fileSelect"></select>
                </div>
            </div>
            <div class="results-content">
                <div class="placeholder" id="placeholder">
                    Upload documents and click "Extract Text" to see results here
                </div>
                <div class="loading hidden" id="loading">
                    <div class="spinner"></div>
                    <span>Extracting text...</span>
                </div>
                <div class="error hidden" id="error"></div>
                <pre class="text-output hidden" id="textOutput"></pre>
            </div>
        </section>

        <footer>
            Uses pdfplumber for PDF and python-docx for Word document processing
        </footer>
    </div>

    <script>
        const dropZone = document.getElementById('dropZone');
        const fileInput = document.getElementById('fileInput');
        const fileList = document.getElementById('fileList');
        const extractBtn = document.getElementById('extractBtn');
        const fileSelector = document.getElementById('fileSelector');
        const fileSelect = document.getElementById('fileSelect');
        const placeholder = document.getElementById('placeholder');
        const loading = document.getElementById('loading');
        const error = document.getElementById('error');
        const textOutput = document.getElementById('textOutput');

        let selectedFiles = [];
        let extractedResults = {};

        // Drop zone events
        dropZone.addEventListener('click', () => fileInput.click());

        dropZone.addEventListener('dragover', (e) => {
            e.preventDefault();
            dropZone.classList.add('drag-over');
        });

        dropZone.addEventListener('dragleave', () => {
            dropZone.classList.remove('drag-over');
        });

        dropZone.addEventListener('drop', (e) => {
            e.preventDefault();
            dropZone.classList.remove('drag-over');
            handleFiles(e.dataTransfer.files);
        });

        fileInput.addEventListener('change', (e) => {
            handleFiles(e.target.files);
        });

        function handleFiles(files) {
            const validExtensions = ['pdf', 'docx', 'doc'];
            
            for (const file of files) {
                const ext = file.name.split('.').pop().toLowerCase();
                if (validExtensions.includes(ext)) {
                    if (!selectedFiles.some(f => f.name === file.name)) {
                        selectedFiles.push(file);
                    }
                }
            }
            
            updateFileList();
            extractBtn.disabled = selectedFiles.length === 0;
        }

        function updateFileList() {
            fileList.innerHTML = selectedFiles.map((file, index) => `
                <div class="file-item">
                    <span class="file-item-name">
                        <svg class="file-item-icon" width="16" height="16" fill="none" stroke="currentColor" viewBox="0 0 24 24">
                            <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M9 12h6m-6 4h6m2 5H7a2 2 0 01-2-2V5a2 2 0 012-2h5.586a1 1 0 01.707.293l5.414 5.414a1 1 0 01.293.707V19a2 2 0 01-2 2z"/>
                        </svg>
                        ${file.name}
                    </span>
                    <button class="file-item-remove" onclick="removeFile(${index})">&times;</button>
                </div>
            `).join('');
        }

        function removeFile(index) {
            selectedFiles.splice(index, 1);
            updateFileList();
            extractBtn.disabled = selectedFiles.length === 0;
        }

        extractBtn.addEventListener('click', async () => {
            if (selectedFiles.length === 0) return;

            // Show loading
            placeholder.classList.add('hidden');
            error.classList.add('hidden');
            textOutput.classList.add('hidden');
            loading.classList.remove('hidden');
            extractBtn.disabled = true;

            const formData = new FormData();
            selectedFiles.forEach(file => {
                formData.append('files', file);
            });

            try {
                const response = await fetch('/extract', {
                    method: 'POST',
                    body: formData
                });

                const data = await response.json();

                if (data.error) {
                    throw new Error(data.error);
                }

                extractedResults = data.results;
                
                // Update file selector
                const fileNames = Object.keys(extractedResults);
                
                if (fileNames.length > 1) {
                    fileSelect.innerHTML = fileNames.map(name => 
                        `<option value="${name}">${name}</option>`
                    ).join('');
                    fileSelector.classList.remove('hidden');
                } else {
                    fileSelector.classList.add('hidden');
                }

                // Show first result
                if (fileNames.length > 0) {
                    displayResult(fileNames[0]);
                }

            } catch (err) {
                error.textContent = err.message || 'An error occurred while extracting text.';
                error.classList.remove('hidden');
            } finally {
                loading.classList.add('hidden');
                extractBtn.disabled = false;
            }
        });

        fileSelect.addEventListener('change', (e) => {
            displayResult(e.target.value);
        });

        function displayResult(fileName) {
            const text = extractedResults[fileName];
            if (text) {
                textOutput.textContent = text;
                textOutput.classList.remove('hidden');
                placeholder.classList.add('hidden');
                error.classList.add('hidden');
            }
        }
    </script>
</body>
</html>
'''


@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)


@app.route('/extract', methods=['POST'])
def extract():
    if 'files' not in request.files:
        return jsonify({'error': 'No files uploaded'}), 400
    
    files = request.files.getlist('files')
    
    if not files or all(f.filename == '' for f in files):
        return jsonify({'error': 'No files selected'}), 400
    
    results = {}
    
    for file in files:
        if file and file.filename and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                # Extract text
                extracted_text = extract_text(tmp_path, filename)
                results[filename] = extracted_text
            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
    
    if not results:
        return jsonify({'error': 'No valid files to process'}), 400
    
    return jsonify({'results': results})


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
